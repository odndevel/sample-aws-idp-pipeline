"""DOCX rendering utilities using python-docx."""

import io
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _ensure_styles(doc: Document) -> None:
    """Ensure custom styles exist in the document."""
    styles = doc.styles

    # Check if styles already exist
    style_names = [s.name for s in styles]

    if "Heading 1" not in style_names:
        h1 = styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
        h1.font.size = Pt(24)
        h1.font.bold = True

    if "Heading 2" not in style_names:
        h2 = styles.add_style("Heading 2", WD_STYLE_TYPE.PARAGRAPH)
        h2.font.size = Pt(20)
        h2.font.bold = True

    if "Heading 3" not in style_names:
        h3 = styles.add_style("Heading 3", WD_STYLE_TYPE.PARAGRAPH)
        h3.font.size = Pt(16)
        h3.font.bold = True


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text with bold and italic formatting to a paragraph.

    Handles **bold** and *italic* markdown syntax.
    """
    # Pattern to match **bold** and *italic*
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*)"

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before the match
        if match.start() > last_end:
            paragraph.add_run(text[last_end : match.start()])

        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def text_to_docx(content: str) -> bytes:
    """Convert plain text to DOCX."""
    doc = Document()

    paragraphs = content.split("\n\n")

    for para in paragraphs:
        if para.strip():
            # Handle single newlines as line breaks
            lines = para.split("\n")
            p = doc.add_paragraph()
            for i, line in enumerate(lines):
                p.add_run(line)
                if i < len(lines) - 1:
                    p.add_run().add_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def markdown_to_docx(content: str) -> bytes:
    """Convert markdown to DOCX with formatting."""
    doc = Document()

    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # Headings
        if line.startswith("#### "):
            p = doc.add_paragraph(line[5:], style="Heading 4")
        elif line.startswith("### "):
            p = doc.add_paragraph(line[4:], style="Heading 3")
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:], style="Heading 2")
        elif line.startswith("# "):
            p = doc.add_paragraph(line[2:], style="Heading 1")
        # Horizontal rule
        elif line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        # Unordered list
        elif line.startswith("- ") or line.startswith("* "):
            text = line[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
        # Ordered list
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
        # Normal paragraph with formatting
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_docx(content: str, format: str) -> bytes:
    """Render content to DOCX bytes.

    Args:
        content: The text or markdown content
        format: "text" or "markdown"

    Returns:
        DOCX bytes
    """
    if format == "markdown":
        return markdown_to_docx(content)
    return text_to_docx(content)
